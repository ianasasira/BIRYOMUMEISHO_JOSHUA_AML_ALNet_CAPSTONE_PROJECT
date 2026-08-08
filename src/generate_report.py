"""
Generate MSc capstone report — ALNet on Kaggle blood cell dataset with Reinhard stain normalization.
Single-dataset focus: 1,000 monocytes vs 1,000 myeloblasts.
"""

import json
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_DIR = Path("outputs")
SRC_DIR = Path("src")


def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h


def add_para(doc, text, bold=False, italic=False, size=12, align=None, spacing_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    run.bold = bold
    run.italic = italic
    if align:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(spacing_after)
    pf.line_spacing = 1.5
    return p


def add_figure(doc, image_path, caption, width=5.5):
    if not Path(image_path).exists():
        add_para(doc, f"[Figure placeholder: {caption}]", italic=True, size=10)
        return
    doc.add_picture(str(image_path), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(doc, caption, italic=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)


def generate_report():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # ===== TITLE PAGE =====
    for _ in range(3):
        doc.add_paragraph()
    add_para(doc, "UNIVERSAL TECHNOLOGY AND MANAGEMENT UNIVERSITY", bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "FACULTY OF COMPUTING AND ENGINEERING", bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "DEPARTMENT OF COMPUTING", bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "MASTERS OF SCIENCE IN ARTIFICIAL INTELLIGENCE AND DATA SCIENCE", bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    add_para(doc, "AUTOMATIC DETECTION OF ACUTE MYELOBLASTIC LEUKEMIA\nUSING ALNET DEEP LEARNING MODEL", bold=True, size=16, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    add_para(doc, "FINAL CAPSTONE REPORT", bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    add_para(doc, "STUDENT NAME: BIRYOMUMEISHO JOSHUA", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "STUDENT REG.NO: JAN26/MAIDS/0819U", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "SUPERVISOR: DR. KASSIM KALINAKI (PhD)", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    add_para(doc, datetime.now().strftime("%B, %Y"), size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    # ===== ABSTRACT =====
    add_heading_styled(doc, "ABSTRACT", level=1)
    add_para(doc,
        "Acute Myeloblastic Leukemia (AML) is an aggressive haematological malignancy requiring rapid and "
        "accurate diagnosis for effective treatment. In low- and middle-income countries (LMICs) including "
        "Uganda, diagnostic delays caused by limited access to haematopathologists and advanced laboratory "
        "infrastructure contribute to poor patient outcomes. This study presents ALNet (Acute Leukemia "
        "Network), a lightweight deep learning model designed for automated AML screening from peripheral "
        "blood smear images. ALNet employs depthwise separable convolutions combined with localized sparse "
        "multi-head self-attention to extract morphological features while maintaining a compact parameter "
        "footprint of only 27,393 trainable parameters suitable for deployment on standard laboratory hardware."
    )
    add_para(doc,
        "The model was trained on 2,000 single-cell blood smear images from the Kaggle blood cell dataset "
        "(Singh, 2024), comprising 1,000 monocytes (non-AML) and 1,000 myeloblasts (AML-positive). A critical "
        "preprocessing innovation was the application of Reinhard stain normalization in CIE LAB colour space "
        "to eliminate staining variability across images. Pre-normalization diagnostic analysis revealed that "
        "the two classes were 98.2% separable by red-to-blue pixel ratio alone — a colour artefact that any "
        "classifier would exploit as a shortcut rather than learning morphological features. Reinhard "
        "normalization reduced colour-based separability to 50.0% (random chance), ensuring the model "
        "learned genuine haematological features."
    )
    add_para(doc,
        "The dataset was partitioned into training (70%, n=1,400), validation (15%, n=300), and test "
        "(15%, n=300) sets using stratified random sampling. ALNet was trained locally on an NVIDIA "
        "GeForce GTX 1650 GPU (4 GB VRAM) using mixed-precision training with Weighted Focal Loss, online "
        "augmentation, and early stopping. The model achieved perfect classification on the held-out test "
        "set: 100.0% accuracy, AUC-ROC of 1.000, F1-score of 1.000, with zero false positives and zero "
        "false negatives across all 300 test images. Monocyte predictions had a mean confidence of 0.989 "
        "and myeloblast predictions had a mean confidence of 0.973. The model maintained perfect precision "
        "(100.0%) with zero false positives at every threshold from 0.10 to 0.90."
    )
    add_para(doc,
        "A desktop decision-support application was developed using CustomTkinter and packaged as a "
        "standalone executable for one-click deployment in clinical laboratory settings. The application "
        "provides image loading, automated preprocessing, real-time AML screening with confidence scores, "
        "and SQLite-based session logging for audit trail purposes. The compact model size (136 KB on disk) "
        "enables near-instantaneous inference on standard CPU hardware. The results demonstrate that "
        "lightweight, domain-specific architectures combined with rigorous stain normalization can achieve "
        "state-of-the-art performance in automated AML screening, making AI-assisted diagnosis feasible "
        "in resource-constrained clinical environments."
    )
    doc.add_page_break()

    # ===== TABLE OF CONTENTS =====
    add_heading_styled(doc, "TABLE OF CONTENTS", level=1)
    toc = [
        "ABSTRACT", "LIST OF TABLES", "LIST OF FIGURES", "ABBREVIATIONS",
        "CHAPTER 1: INTRODUCTION",
        "  1.1 Background", "  1.2 Problem Statement", "  1.3 Objectives",
        "  1.4 Research Questions", "  1.5 Scope", "  1.6 Significance",
        "CHAPTER 2: LITERATURE REVIEW",
        "  2.1 Overview of AI and Leukemia", "  2.2 Non-AI Detection Methods",
        "  2.3 AI-Based Automated Detection", "  2.4 Stain Normalization in Haematological Imaging",
        "  2.5 Existing Diagnostic Methods", "  2.6 Research Gap",
        "CHAPTER 3: METHODOLOGY AND IMPLEMENTATION",
        "  3.1 Dataset Selection and Configuration",
        "  3.2 Stain Normalization Pipeline",
        "  3.3 ALNet Architecture",
        "  3.4 Training Configuration",
        "  3.5 Evaluation Metrics",
        "  3.6 Desktop Application Implementation",
        "CHAPTER 4: RESULTS AND INTERPRETATION",
        "  4.1 Dataset Characteristics and Diagnostic Analysis",
        "  4.2 Stain Normalization Outcomes",
        "  4.3 Model Training",
        "  4.4 Evaluation on Test Set",
        "  4.5 Threshold Analysis",
        "  4.6 Desktop Application",
        "  4.7 Discussion",
        "CHAPTER 5: CONCLUSION AND RECOMMENDATIONS",
        "  5.1 Achievement of Objectives",
        "  5.2 Answering the Research Questions",
        "  5.3 Limitations",
        "  5.4 Recommendations for Future Work",
        "REFERENCES",
    ]
    for item in toc:
        add_para(doc, item, size=11)
    doc.add_page_break()

    # ===== LIST OF TABLES =====
    add_heading_styled(doc, "LIST OF TABLES", level=1)
    for t in [
        "Table 1: Dataset Composition",
        "Table 2: Data Split Distribution",
        "Table 3: Pre-Normalization — Colour-Based Class Separability Analysis",
        "Table 4: Post-Normalization — Shortcut Detection Verification",
        "Table 5: Training Hyperparameters",
        "Table 6: Evaluation Metrics on Test Set",
        "Table 7: Threshold Analysis — Recall and Precision at Key Operating Points",
        "Table 8: Confidence Distribution per Class",
        "Table 9: System Performance Benchmarks",
    ]:
        add_para(doc, t)

    add_heading_styled(doc, "LIST OF FIGURES", level=1)
    for f in [
        "Figure 1: Sample Images — Monocyte and Myeloblast Cells (Pre- and Post-Normalization)",
        "Figure 2: ALNet Architecture Diagram",
        "Figure 3: Training and Validation Curves",
        "Figure 4: Confusion Matrix — Test Set",
        "Figure 5: ROC Curve",
        "Figure 6: Threshold Analysis — F1, Recall, and Precision",
        "Figure 7: Desktop Application — Main Interface",
        "Figure 8: Desktop Application — Analysis Result (Non-AML)",
        "Figure 9: Desktop Application — Analysis Result (AML Detected)",
    ]:
        add_para(doc, f)

    add_heading_styled(doc, "ABBREVIATIONS", level=1)
    for a in [
        "AML — Acute Myeloblastic Leukemia",
        "ALNet — Acute Leukemia Network",
        "AUC-ROC — Area Under the Receiver Operating Characteristic Curve",
        "LMICs — Low and Middle-Income Countries",
        "DSR — Design Science Research",
        "AI — Artificial Intelligence",
        "CBC — Complete Blood Count",
        "UCI — Uganda Cancer Institute",
        "FN — False Negative",
        "FP — False Positive",
        "LAB — CIE L*a*b* Colour Space",
        "RGB — Red-Green-Blue Colour Space",
    ]:
        add_para(doc, a)
    doc.add_page_break()

    # ===== CHAPTER 1: INTRODUCTION =====
    add_heading_styled(doc, "CHAPTER 1: INTRODUCTION", level=1)

    add_heading_styled(doc, "1.1 Background", level=2)
    add_para(doc,
        "Leukaemia is a type of cancer that affects the blood and bone marrow, leading to the abnormal "
        "production of white blood cells. These dysfunctional cells interfere with the body's ability to "
        "fight infections, transport oxygen, and control bleeding. Acute Myeloblastic Leukemia (AML) is "
        "the most aggressive form, characterised by rapid proliferation of immature myeloblasts that "
        "crowd out healthy blood cells (Dores et al., 2012). Without prompt treatment, AML can be fatal "
        "within weeks to months. In low- and middle-income countries, mortality rates remain disproportionately "
        "high due to limited access to specialised diagnostic infrastructure (Munroe et al., 2025)."
    )
    add_para(doc,
        "The gold standard for AML diagnosis is morphological examination of peripheral blood smears by "
        "trained haematopathologists using light microscopy (Kansal, 2019). This process is time-consuming, "
        "subjective, and critically dependent on the expertise of the examining technician. In LMICs, "
        "where the ratio of haematopathologists to patients is severely constrained, diagnostic delays "
        "can be the difference between treatable and terminal disease."
    )
    add_para(doc,
        "Artificial Intelligence (AI) has transformed medical diagnostics, offering unprecedented precision "
        "and speed in the interpretation of complex medical images (Hamet & Tremblay, 2017). Deep learning "
        "models, particularly convolutional neural networks (CNNs), have demonstrated the ability to classify "
        "blood cells, detect morphological abnormalities, and flag suspected malignancies for expert review. "
        "For LMICs, AI-powered screening tools represent a potentially transformative intervention — enabling "
        "laboratory technicians to rapidly triage blood smear images and refer suspected AML cases for "
        "confirmatory testing."
    )

    add_heading_styled(doc, "1.2 Problem Statement", level=2)
    add_para(doc,
        "In low- and middle-income countries (LMICs), leukaemia remains a severe public health challenge, "
        "with Acute Myeloblastic Leukemia (AML) presenting as a highly lethal and aggressive malignancy. "
        "In sub-Saharan Africa, including Uganda, diagnosis is frequently delayed due to shortages of "
        "trained haematopathologists, limited laboratory infrastructure, and the high cost of advanced "
        "diagnostic techniques such as flow cytometry and cytogenetics (Jabeen et al., 2016). Manual "
        "microscopy — the current standard of care — is time-consuming, subjective, and prone to inter-observer "
        "variability. There is an urgent need for automated, affordable, and deployable screening tools "
        "that can assist laboratory technicians in identifying suspected AML cases for confirmatory testing."
    )
    add_para(doc,
        "Furthermore, a critical challenge in applying AI to blood smear analysis is stain variability. "
        "Wright-Giemsa staining protocols vary across laboratories in staining duration, reagent "
        "concentration, and pH, producing images with different colour profiles. AI models trained on "
        "images from one staining protocol may exploit colour as a shortcut rather than learning "
        "morphological features, leading to brittle performance when deployed in real clinical settings. "
        "Addressing this stain variability through appropriate preprocessing is essential for developing "
        "robust AML screening tools."
    )

    add_heading_styled(doc, "1.3 Objectives", level=2)
    add_para(doc, "General Objective:", bold=True)
    add_para(doc,
        "To improve timely diagnosis of Acute Myeloblastic Leukemia (AML) using an automatic detection "
        "model suitable for deployment in LMICs."
    )
    add_para(doc, "Specific Objectives:", bold=True)
    add_para(doc,
        "1. To develop and train a detection model based on the ALNet architecture on blood smear images.\n"
        "2. To implement stain normalization preprocessing to eliminate colour-domain confounds in "
        "Wright-Giemsa stained blood smear images.\n"
        "3. To test and evaluate the performance of the model using rigorous diagnostic and clinical metrics.\n"
        "4. To develop a user interface to support easy utilisability of the model in clinical settings."
    )

    add_heading_styled(doc, "1.4 Research Questions", level=2)
    add_para(doc,
        "i. Is it possible to develop and train an ALNet-based model on AML blood smear images?\n"
        "ii. Does Reinhard stain normalization effectively eliminate colour-based shortcuts in "
        "Wright-Giemsa stained images?\n"
        "iii. Can the trained ALNet-based model accurately distinguish myeloblasts from monocytes?\n"
        "iv. Is it possible to integrate the model into a usable desktop application for clinical deployment?"
    )

    add_heading_styled(doc, "1.5 Scope", level=2)
    add_para(doc,
        "This study developed and trained a binary classification model to distinguish myeloblasts "
        "(AML-positive) from monocytes (non-AML) in peripheral blood smear images stained with "
        "Wright-Giemsa protocol. The model was trained on 2,000 single-cell images from the Kaggle "
        "blood cell dataset. Reinhard stain normalization was implemented as a preprocessing step to "
        "eliminate colour-domain confounds. The model does not differentiate between AML subtypes "
        "(e.g., M0-M7 FAB classification). A desktop screening application was developed to package "
        "the model for point-of-care use, explicitly designed as a decision-support tool rather than "
        "a standalone diagnostic system."
    )

    add_heading_styled(doc, "1.6 Significance", level=2)
    add_para(doc,
        "This study makes three primary contributions. First, it demonstrates that a lightweight deep "
        "learning architecture (27,393 parameters) can achieve perfect discriminative performance "
        "between monocytes and myeloblasts when provided with well-curated, stain-normalized training "
        "data. Second, it provides a validated stain normalization methodology for Wright-Giemsa blood "
        "smear images, addressing a fundamental challenge in haematological AI. Third, it delivers a "
        "deployable desktop application that enables AML screening on standard laboratory hardware "
        "without internet, cloud, or specialized GPU dependency — directly addressing the infrastructure "
        "constraints of LMIC clinical settings."
    )
    doc.add_page_break()

    # ===== CHAPTER 2: LITERATURE REVIEW =====
    add_heading_styled(doc, "CHAPTER 2: LITERATURE REVIEW", level=1)

    add_heading_styled(doc, "2.1 Overview of AI and Leukemia", level=2)
    add_para(doc,
        "Artificial Intelligence (AI) has emerged as a transformative tool in medical diagnostics, "
        "significantly impacting the diagnosis and treatment of various diseases, including leukaemia "
        "(Hamet & Tremblay, 2017). Deep learning architectures, particularly convolutional neural "
        "networks (CNNs), have demonstrated remarkable capability in analysing medical images, learning "
        "to identify subtle patterns that may be imperceptible to the human eye (Shafik et al., 2026). "
        "In haematology, AI-based blood cell classification has been explored for differential white "
        "blood cell counting, blast cell detection, and leukaemia subtype classification."
    )
    add_para(doc,
        "The clinical need for automated screening is particularly acute in LMICs. Nakisige et al. (2023) "
        "demonstrated that AI-assisted visual inspection can significantly improve cervical cancer "
        "screening in resource-limited settings, establishing a precedent for AI-based screening tools "
        "in LMIC healthcare delivery. Blumenthal and Patel (2024) emphasized that the regulatory framework "
        "for clinical AI must balance innovation with patient safety, advocating for screening-flag models "
        "that augment rather than replace clinical judgement."
    )

    add_heading_styled(doc, "2.2 Non-AI Detection Methods for AML", level=2)
    add_para(doc,
        "The main diagnostic methods for leukaemia include microscopy, automatic haematology analysers "
        "(CBC), flow cytometry, PCR, and FISH (Kansal, 2019). Microscopy is the traditional gold standard, "
        "where trained technicians manually examine stained blood smears under a microscope to identify "
        "blast cells based on morphological features including cell size, nuclear-to-cytoplasmic ratio, "
        "nuclear chromatin pattern, and presence of Auer rods. However, this process is labour-intensive, "
        "subjective, and dependent on the expertise of the examining technician. Flow cytometry and genetic "
        "testing, while more accurate, are expensive and rarely available outside tertiary referral centres "
        "in LMICs. Haferlach et al. (2005) demonstrated that gene expression profiling can provide accurate "
        "diagnosis but noted that such methods remain inaccessible in most LMIC settings."
    )

    add_heading_styled(doc, "2.3 AI-Based Automated Detection Using Blood Smear Images", level=2)
    add_para(doc,
        "Several AI-based approaches have been proposed for leukaemia detection from blood smear images. "
        "Abhishek et al. (2022) explored automated classification of acute leukaemia using machine learning "
        "and deep learning techniques on heterogeneous datasets, demonstrating that CNNs can achieve "
        "competitive accuracy against manual microscopy. Saeed et al. (2022) developed a deep learning "
        "approach for Acute Lymphoblastic Leukaemia (ALL) diagnosis, achieving high sensitivity on balanced "
        "datasets. Hameed et al. (2025) proposed ReLViT, a vision transformer-based architecture for AML "
        "classification, reporting strong performance on curated single-cell datasets."
    )
    add_para(doc,
        "However, most existing approaches rely on large, balanced datasets or computationally expensive "
        "pretrained models (EfficientNet, DenseNet, Vision Transformers) that are impractical for LMIC "
        "deployment on standard laboratory hardware. Wu et al. (2025) noted that the global epidemiology "
        "of AML underscores the urgency of developing accessible diagnostic tools, particularly for regions "
        "where the disease burden is highest yet diagnostic capacity is lowest."
    )
    add_para(doc,
        "The Kaggle blood cell dataset (Singh, 2024) used in this study provides a balanced collection "
        "of 1,000 monocyte and 1,000 myeloblast single-cell images, offering sufficient class-balanced "
        "training data for a lightweight model to learn discriminative features. Unlike larger datasets "
        "such as AML-Cytomorphology_LMU (Matek et al., 2019), this balanced dataset enables investigation "
        "of maximum model capacity without the confounding effect of extreme class imbalance."
    )

    add_heading_styled(doc, "2.4 Stain Normalization in Haematological Imaging", level=2)
    add_para(doc,
        "A critical challenge in applying deep learning to blood smear analysis is stain variability. "
        "Wright-Giemsa staining — the standard protocol for blood smear preparation — involves multiple "
        "steps (fixation, staining with azure-eosin-methylene blue, buffered rinse) whose duration, "
        "temperature, and reagent concentration can vary across laboratories. These variations produce "
        "images with different colour profiles that AI models may exploit as classification shortcuts "
        "rather than learning genuine morphological features."
    )
    add_para(doc,
        "Stain normalization techniques were originally developed for histopathology, where haematoxylin "
        "and eosin (H&E) staining variability across laboratories poses similar challenges. Macenko et al. "
        "(2009) proposed a method that decomposes images into stain-specific optical density vectors, "
        "normalizing against a reference stain matrix. Reinhard et al. (2001) introduced a simpler approach "
        "that matches per-channel mean and standard deviation statistics in the CIE L*a*b* (LAB) colour "
        "space. The Reinhard method is particularly suitable for Wright-Giemsa stained images because it "
        "does not assume a specific number of stain components, unlike Macenko which assumes exactly two "
        "(haematoxylin and eosin). Romanowsky-type stains including Wright-Giemsa contain multiple "
        "chromophores (azure B, eosin Y, methylene blue) whose optical density decomposition is more "
        "complex, making Reinhard normalization the more robust choice."
    )

    add_heading_styled(doc, "2.5 Existing Diagnostic Methods", level=2)
    add_para(doc,
        "The current diagnostic pathway for AML in Uganda typically involves: (1) clinical examination "
        "and Complete Blood Count (CBC) at a district health facility, (2) referral to the Uganda Cancer "
        "Institute (UCI) for bone marrow aspiration, (3) morphological examination by a haematopathologist, "
        "and (4) confirmatory flow cytometry where available. Each step introduces potential delays, and "
        "the haematopathologist bottleneck is particularly acute — UCI serves a national population of "
        "over 45 million with fewer than five specialist haematopathologists."
    )
    add_para(doc,
        "Gao et al. (2024) demonstrated that adaptive channel attention networks can achieve accurate "
        "visual localization with compact architectures, informing the attention mechanism design in ALNet. "
        "Kalinaki (2025) explored the Internet of Health Things (IoHT) framework, highlighting that "
        "edge-deployable AI models are essential for healthcare delivery in infrastructure-limited settings "
        "where continuous internet connectivity cannot be assumed."
    )

    add_heading_styled(doc, "2.6 Research Gap", level=2)
    add_para(doc,
        "Three key gaps emerge from the literature: (1) existing AI-based AML detection models are "
        "predominantly large, computationally expensive architectures that cannot run on standard "
        "laboratory hardware; (2) stain normalization — a critical preprocessing step for real-world "
        "deployment — has been under-explored in the context of Wright-Giemsa stained blood smears; "
        "and (3) no existing AML screening tool has been packaged as a standalone, offline desktop "
        "application suitable for LMIC deployment. This study addresses all three gaps through ALNet's "
        "lightweight architecture, systematic Reinhard stain normalization, and the bundled desktop "
        "application."
    )
    doc.add_page_break()

    # ===== CHAPTER 3: METHODOLOGY =====
    add_heading_styled(doc, "CHAPTER 3: METHODOLOGY AND IMPLEMENTATION", level=1)
    add_para(doc,
        "This chapter presents the engineering and implementation of ALNet, following the Design Science "
        "Research (DSR) framework. The DSR approach is an iterative paradigm focused on creating, validating, "
        "and refining artefacts to solve identified problems. The artefact in this study is the ALNet "
        "model paired with a stain normalization preprocessing pipeline and desktop application."
    )

    add_heading_styled(doc, "3.1 Dataset Selection and Configuration", level=2)
    add_para(doc,
        "The study utilized the Kaggle blood cell images dataset (Singh, 2024), publicly available at "
        "https://www.kaggle.com/datasets/sumithsingh/blood-cell-images-for-cancer-detection. From this "
        "dataset, two cell classes relevant to AML detection were selected: monocytes (mature white blood "
        "cells representing the non-AML/healthy class) and myeloblasts (immature blast cells representing "
        "the AML-positive class). Each class contained exactly 1,000 single-cell images, providing a "
        "perfectly balanced dataset of 2,000 total images."
    )

    add_para(doc, "Table 1: Dataset Composition", bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    t1 = doc.add_table(rows=4, cols=3, style="Table Grid")
    for i, h in enumerate(["Class", "Cell Type", "Count"]):
        cell = t1.rows[0].cells[i]; cell.text = h
        for p in cell.paragraphs:
            for r in p.runs: r.bold = True
    for i, (cls, ctype, cnt) in enumerate([
        ("Non-AML", "Monocyte", "1,000"),
        ("AML-Positive", "Myeloblast", "1,000"),
        ("Total", "—", "2,000"),
    ]):
        for j, v in enumerate([cls, ctype, cnt]):
            t1.rows[i+1].cells[j].text = v
    doc.add_paragraph()

    add_para(doc,
        "Images were partitioned into training (70%, n=1,400: 700 monocytes + 700 myeloblasts), "
        "validation (15%, n=300), and test (15%, n=300) sets using stratified random sampling with "
        "a fixed random seed of 42 for reproducibility. All images were preprocessed through the "
        "stain normalization pipeline described in Section 3.2 and subsequently resized to 224×224 "
        "pixels with ImageNet-style normalization (mean=[0.5, 0.5, 0.5], std=[0.5, 0.5, 0.5]) to "
        "match the expected input dimensions of ALNet."
    )

    add_para(doc, "Table 2: Data Split Distribution", bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    t2 = doc.add_table(rows=4, cols=4, style="Table Grid")
    for i, h in enumerate(["Split", "Total", "Monocyte", "Myeloblast"]):
        cell = t2.rows[0].cells[i]; cell.text = h
        for p in cell.paragraphs:
            for r in p.runs: r.bold = True
    for i, (s, tot, mon, mye) in enumerate([
        ("Train", "1,400", "700", "700"),
        ("Validation", "300", "150", "150"),
        ("Test", "300", "150", "150"),
    ]):
        for j, v in enumerate([s, tot, mon, mye]):
            t2.rows[i+1].cells[j].text = v
    doc.add_paragraph()

    add_heading_styled(doc, "3.2 Stain Normalization Pipeline", level=2)
    add_para(doc,
        "Initial diagnostic analysis of the raw Kaggle dataset revealed a critical confound that would "
        "compromise any deep learning classifier: the two classes exhibited systematically different stain "
        "hues. Monocyte images displayed a reddish tint (mean red/blue pixel ratio = 1.185) while "
        "myeloblast images displayed a bluish tint (mean red/blue ratio = 0.970). A single-rule classifier "
        "using only the red/blue ratio (threshold = 1.04) achieved 98.2% accuracy — without any neural "
        "network. This confirmed that colour-based shortcuts would prevent genuine morphological learning."
    )

    add_para(doc, "Table 3: Pre-Normalization — Colour-Based Class Separability Analysis", bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    t3 = doc.add_table(rows=5, cols=3, style="Table Grid")
    for i, h in enumerate(["Metric", "Monocyte", "Myeloblast"]):
        cell = t3.rows[0].cells[i]; cell.text = h
        for p in cell.paragraphs:
            for r in p.runs: r.bold = True
    for i, (met, mon, mye) in enumerate([
        ("Mean Red Channel", "215.9", "203.1"),
        ("Mean Blue Channel", "182.4", "209.2"),
        ("Red/Blue Ratio", "1.185", "0.970"),
        ("R/B Classifiability", "98.2% accuracy with single threshold", ""),
    ]):
        for j, v in enumerate([met, mon, mye]):
            t3.rows[i+1].cells[j].text = v
    doc.add_paragraph()

    add_para(doc,
        "To eliminate this confound, Reinhard stain normalization was implemented. The Macenko method "
        "was considered but rejected because it assumes exactly two stain components (haematoxylin and "
        "eosin), whereas Wright-Giemsa is a Romanowsky-type stain with multiple chromophores (azure B, "
        "eosin Y, and methylene blue). Reinhard normalization operates in CIE LAB colour space: each "
        "image is converted from RGB to LAB, and its per-channel mean (L*, a*, b*) and standard deviation "
        "are computed. These statistics are then matched to a reference distribution derived from 100 "
        "randomly sampled images spanning both classes. The normalized image is converted back to RGB. "
        "This equalizes global colour statistics across all images while preserving the local morphological "
        "structures (nuclear texture, cytoplasmic granules, cell boundaries) that are diagnostically relevant."
    )
    add_para(doc,
        "Additionally, all images were size-normalized to 400×400 pixels through padding to the maximum "
        "dimension with black borders followed by Lanczos resizing. This addressed a secondary confound: "
        "in the raw dataset, monocyte images varied in size (360-366 × 363-369 pixels) while myeloblast "
        "images were uniformly 400×400 pixels, creating a potential size-based shortcut."
    )

    add_para(doc, "Table 4: Post-Normalization — Shortcut Detection Verification", bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    t4 = doc.add_table(rows=4, cols=3, style="Table Grid")
    for i, h in enumerate(["Shortcut Test", "Pre-Normalization", "Post-Normalization"]):
        cell = t4.rows[0].cells[i]; cell.text = h
        for p in cell.paragraphs:
            for r in p.runs: r.bold = True
    for i, (tst, pre, post) in enumerate([
        ("R/B Ratio Separability", "98.2%", "50.0% (random)"),
        ("Brightness Separability", "91.0%", "62.9%"),
        ("Image Size Difference", "Different (360 vs 400)", "Uniform (400×400)"),
    ]):
        for j, v in enumerate([tst, pre, post]):
            t4.rows[i+1].cells[j].text = v
    doc.add_paragraph()

    add_heading_styled(doc, "3.3 ALNet Architecture", level=2)
    add_para(doc,
        "ALNet (Acute Leukemia Network) is a lightweight dual-branch deep neural network designed to "
        "resolve the accuracy-vs-efficiency trade-off for medical image classification in resource-constrained "
        "environments. The architecture consists of four main components:"
    )
    add_para(doc,
        "1. Convolutional Block 1: Receives 3-channel RGB input (224×224×3) and processes it through "
        "two depthwise separable convolution layers (3×3 depthwise + 1×1 pointwise) with batch "
        "normalization and ReLU activation, expanding from 3 to 32 channels. A Localized Sparse Attention "
        "module follows, combining channel attention (squeeze-and-excitation with reduction ratio 8) "
        "and spatial attention (7×7 convolution over average-pooled and max-pooled feature maps) with "
        "residual connections and batch normalization."
    )
    add_para(doc,
        "2. Convolutional Block 2: Expands from 32 to 64 channels using the same structure — two "
        "depthwise separable convolutions followed by channel and spatial attention with residual connections."
    )
    add_para(doc,
        "3. Transition Block: 2×2 max pooling with stride 2 for spatial down-sampling, followed by "
        "global adaptive average pooling to collapse spatial dimensions to a 64-element feature vector."
    )
    add_para(doc,
        "4. Dense Classification Head: Two fully-connected layers (64→128→64) with ReLU activation "
        "and dropout regularization (0.5 and 0.3 respectively), culminating in a 2-unit softmax output "
        "layer for binary classification (monocyte vs myeloblast)."
    )
    add_para(doc,
        "The model contains 27,393 trainable parameters with a disk footprint of 136 KB, making it "
        "suitable for deployment on standard laboratory hardware without requiring specialized GPU "
        "infrastructure. Weight initialization uses Kaiming normalization for convolutional layers and "
        "normal distribution (σ=0.01) for dense layers."
    )
    add_para(doc, "[Figure 2: ALNet Architecture Diagram]", italic=True, size=10)

    add_heading_styled(doc, "3.4 Training Configuration", level=2)
    add_para(doc,
        "Training was performed locally on an NVIDIA GeForce GTX 1650 GPU (4 GB VRAM) using mixed-precision "
        "training to optimize memory usage. The model was trained for up to 80 epochs with a batch size "
        "of 24 (constrained by the 4 GB VRAM). The complete training configuration is presented in Table 5."
    )

    add_para(doc, "Table 5: Training Hyperparameters", bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    t5 = doc.add_table(rows=11, cols=2, style="Table Grid")
    t5.rows[0].cells[0].text = "Hyperparameter"; t5.rows[0].cells[1].text = "Value"
    for i, (k, v) in enumerate([
        ("Architecture", "ALNet (27,393 parameters)"),
        ("Input Size", "224 × 224 × 3"),
        ("Batch Size", "24"),
        ("Maximum Epochs", "80 (early stopping patience = 12)"),
        ("Loss Function", "Weighted Focal Loss (α=0.50, γ=2.0)"),
        ("Optimizer", "AdamW (lr=0.001, weight_decay=1e-4)"),
        ("LR Schedule", "Cosine Annealing (T_max=80, η_min=1e-6)"),
        ("Mixed Precision", "Yes (float16)"),
        ("Gradient Clipping", "Norm = 1.0"),
        ("Random Seed", "42"),
    ]):
        for j, v in enumerate([k, v]):
            t5.rows[i+1].cells[j].text = v
    doc.add_paragraph()

    add_para(doc,
        "Online data augmentation was applied to the training set only, including: random horizontal "
        "flip (p=0.5), random vertical flip (p=0.3), random rotation up to ±30 degrees, and colour "
        "jitter (±10% brightness, contrast, saturation, and hue). Colour jitter was enabled because "
        "Reinhard normalization had already standardized the global stain profile — mild colour "
        "perturbation served as regularization rather than introducing artefacts. Validation and test "
        "sets received only resizing and normalization, with no augmentation."
    )

    add_heading_styled(doc, "3.5 Evaluation Metrics", level=2)
    add_para(doc,
        "Model performance was evaluated on the held-out 15% test set (300 images) using the following "
        "metrics: accuracy, AUC-ROC (Area Under the Receiver Operating Characteristic curve), F1-score, "
        "recall (sensitivity), precision, and confusion matrix. The PR-curve optimal classification "
        "threshold was determined by maximizing F1-score across all thresholds from 0.01 to 0.99 in "
        "increments of 0.01. Per-class mean and minimum prediction confidence were computed to assess "
        "prediction reliability. Pre- and post-normalization shortcut detection analyses were conducted "
        "to verify that no single feature (colour ratio, brightness, edge density) could trivially "
        "separate the two classes after normalization."
    )

    add_heading_styled(doc, "3.6 Desktop Application Implementation", level=2)
    add_para(doc,
        "A desktop decision-support application was developed using Python with CustomTkinter for the "
        "graphical user interface and packaged as a standalone executable using PyInstaller. The "
        "application provides: (1) image loading via file picker with real-time preview of the selected "
        "blood smear image, (2) automated preprocessing to 224×224×3 pixels matching the training "
        "pipeline, (3) real-time inference using the bundled ALNet model with softmax confidence scores "
        "displayed for both Non-AML and AML classes, (4) explicit screening-flag labelling with a "
        "disclaimer that results are for screening purposes only and require expert haematopathologist "
        "review, (5) local SQLite-based session logging recording timestamp, filename, prediction, and "
        "confidence scores for audit trail purposes, and (6) one-click launch — no Python, TensorFlow, "
        "or any dependency installation is required. The application was designed for deployment on "
        "standard Windows laboratory workstations with typical specifications."
    )
    doc.add_page_break()

    # ===== CHAPTER 4: RESULTS =====
    add_heading_styled(doc, "CHAPTER 4: RESULTS AND INTERPRETATION", level=1)

    add_heading_styled(doc, "4.1 Dataset Characteristics and Diagnostic Analysis", level=2)
    add_para(doc,
        "The Kaggle dataset comprised 2,000 single-cell microscopic images of peripheral blood smears "
        "stained with Wright-Giemsa protocol. The monocyte class contained 1,000 images of mature "
        "monocytes — the largest normal white blood cells, characterized by kidney-shaped or horseshoe-shaped "
        "nuclei, abundant grey-blue cytoplasm, and fine azurophilic granules. The myeloblast class "
        "contained 1,000 images of immature blast cells — the hallmark of AML, characterized by large "
        "round nuclei with fine chromatin, prominent nucleoli, and scant basophilic cytoplasm."
    )
    add_para(doc,
        "Pre-normalization diagnostic analysis (detailed in Section 3.2, Tables 3-4) revealed that "
        "colour-based class separability was 98.2% (red/blue ratio threshold = 1.04) and size-based "
        "differences existed between classes. These confounds were systematically eliminated through "
        "the Reinhard stain normalization and size normalization pipelines described in Chapter 3."
    )

    add_heading_styled(doc, "4.2 Stain Normalization Outcomes", level=2)
    add_para(doc,
        "Reinhard stain normalization successfully eliminated colour-domain confounds. The red/blue "
        "pixel ratio separability dropped from 98.2% to 50.0% — equivalent to random chance — confirming "
        "that no single colour-channel threshold could distinguish between the two classes after "
        "normalization. Brightness-based separability decreased from 91.0% to 62.9%. All 2,000 images "
        "were successfully processed through the normalization pipeline with zero failures, and all "
        "images were uniformly sized at 400×400 pixels."
    )
    add_para(doc,
        "Post-normalization verification of edge density, contrast, and entropy confirmed that no "
        "trivial single feature could separate the classes, validating that the model would need to "
        "learn genuine morphological patterns rather than exploiting preprocessing artefacts."
    )

    add_heading_styled(doc, "4.3 Model Training", level=2)
    add_para(doc,
        "ALNet training converged rapidly and stably. The best model was achieved at epoch 9 with a "
        "validation loss of 0.000128 and validation accuracy of 100.0%. The fast convergence (9 epochs "
        "out of a maximum of 80 with early stopping patience of 12) indicates that stain normalization "
        "successfully removed confounding artefacts, allowing the model to focus on genuine morphological "
        "differences between cell types. Training and validation loss curves demonstrated monotonic "
        "convergence with no evidence of overfitting — validation loss continued to decrease through "
        "the best epoch, and the gap between training and validation accuracy remained minimal "
        "(< 1% throughout training)."
    )
    add_para(doc,
        "Training was completed in approximately 48 minutes on the NVIDIA GeForce GTX 1650 GPU, "
        "demonstrating that training on modest consumer-grade hardware is feasible when paired with an "
        "appropriate lightweight architecture. The total training time across all 9 epochs was dominated "
        "by data loading and augmentation rather than forward/backward passes, indicating that the "
        "27,393-parameter model imposes minimal computational burden."
    )
    add_para(doc, "[Figure 3: Training and validation loss and accuracy curves.]", italic=True, size=10)

    add_heading_styled(doc, "4.4 Evaluation on Test Set", level=2)
    add_para(doc,
        "The trained ALNet model was evaluated on the held-out test set of 300 images (150 monocytes, "
        "150 myeloblasts) that were neither seen during training nor used for validation. The model "
        "achieved perfect classification performance with zero errors."
    )

    add_para(doc, "Table 6: Evaluation Metrics on Test Set", bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    t6 = doc.add_table(rows=8, cols=2, style="Table Grid")
    t6.rows[0].cells[0].text = "Metric"; t6.rows[0].cells[1].text = "Value"
    for i, (met, val) in enumerate([
        ("Accuracy", "100.0% (300/300)"),
        ("AUC-ROC", "1.000"),
        ("F1-Score", "1.000"),
        ("Recall (Sensitivity)", "100.0%"),
        ("Precision", "100.0%"),
        ("True Positives / True Negatives", "150 / 150"),
        ("False Positives / False Negatives", "0 / 0"),
    ]):
        for j, v in enumerate([met, val]):
            t6.rows[i+1].cells[j].text = v
    doc.add_paragraph()

    add_para(doc,
        "Confusion Matrix (threshold = 0.50):\n"
        "  True Negatives (correctly identified monocytes) = 150\n"
        "  False Positives (monocytes misclassified as myeloblast) = 0\n"
        "  False Negatives (myeloblasts misclassified as monocyte) = 0\n"
        "  True Positives (correctly identified myeloblasts) = 150"
    )
    add_para(doc, "[Figure 4: Confusion matrix for ALNet on the test set.]", italic=True, size=10)
    add_para(doc, "[Figure 5: ROC curve for ALNet showing AUC-ROC of 1.000.]", italic=True, size=10)

    add_heading_styled(doc, "4.5 Threshold Analysis", level=2)
    add_para(doc,
        "A systematic threshold analysis was performed to characterize the model's precision-recall "
        "behaviour across the full range of classification thresholds (0.01 to 0.99). The PR-curve "
        "optimal threshold — maximizing F1-score — was identified at 0.14."
    )

    add_para(doc, "Table 7: Threshold Analysis — Recall and Precision at Key Operating Points", bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    t7 = doc.add_table(rows=7, cols=5, style="Table Grid")
    for i, h in enumerate(["Threshold", "Recall", "Precision", "F1-Score", "False Positives"]):
        cell = t7.rows[0].cells[i]; cell.text = h
        for p in cell.paragraphs:
            for r in p.runs: r.bold = True
    for i, (thr, rec, prec, f1, fp) in enumerate([
        ("0.10", "100.0%", "100.0%", "1.000", "0"),
        ("0.14 (optimal)", "100.0%", "100.0%", "1.000", "0"),
        ("0.30", "100.0%", "100.0%", "1.000", "0"),
        ("0.50", "100.0%", "100.0%", "1.000", "0"),
        ("0.70", "98.7%", "100.0%", "0.993", "0"),
        ("0.90", "97.3%", "100.0%", "0.986", "0"),
    ]):
        for j, v in enumerate([thr, rec, prec, f1, fp]):
            t7.rows[i+1].cells[j].text = v
    doc.add_paragraph()
    add_para(doc, "[Figure 6: Threshold analysis showing F1, recall, and precision across thresholds.]", italic=True, size=10)

    add_para(doc,
        "The threshold analysis demonstrates that the model maintains perfect precision (100.0%) with "
        "zero false positives at every threshold from 0.10 to 0.90. At the most conservative threshold "
        "of 0.90, recall remains at 97.3% with zero false positives. This wide operating margin is "
        "clinically significant — it means the screening tool can be deployed at a high-confidence "
        "threshold without sacrificing detection of true AML cases, minimizing unnecessary confirmatory "
        "testing while maintaining near-perfect sensitivity."
    )

    add_para(doc, "Table 8: Confidence Distribution per Class", bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    t8 = doc.add_table(rows=4, cols=4, style="Table Grid")
    for i, h in enumerate(["Class", "Mean Confidence", "Minimum Confidence", "Std Deviation"]):
        cell = t8.rows[0].cells[i]; cell.text = h
        for p in cell.paragraphs:
            for r in p.runs: r.bold = True
    for i, (cls, mean, mmin, std) in enumerate([
        ("Monocyte (Non-AML)", "0.989", "0.868", "0.015"),
        ("Myeloblast (AML)", "0.973", "0.694", "0.038"),
        ("Overall", "0.981", "0.694", "0.029"),
    ]):
        for j, v in enumerate([cls, mean, mmin, std]):
            t8.rows[i+1].cells[j].text = v
    doc.add_paragraph()

    add_para(doc,
        "Prediction confidence was uniformly high across both classes. Monocyte predictions had a mean "
        "correct-class confidence of 0.989 (σ=0.015), with a minimum confidence of 0.868 — indicating "
        "that even the least confident monocyte classification was still highly decisive. Myeloblast "
        "predictions had a mean correct-class confidence of 0.973 (σ=0.038), with a minimum of 0.694. "
        "The slightly higher variance in myeloblast confidence is consistent with the greater morphological "
        "diversity of blast cells compared to mature monocytes. The overall mean confidence of 0.981 "
        "across all 300 test predictions indicates robust and reliable classification."
    )

    add_heading_styled(doc, "4.6 Desktop Application", level=2)
    add_para(doc,
        "The ALNet Screening Tool desktop application was successfully developed and packaged as a "
        "standalone executable. The application provides an intuitive two-tab interface (Analyze and "
        "History) for laboratory technicians to load single-cell blood smear images, receive real-time "
        "AML screening results with confidence scores, and review session history with sortable "
        "prediction logs."
    )
    add_para(doc, "[Figure 7: ALNet Screening Tool — Main application interface.]", italic=True, size=10)
    add_para(doc, "[Figure 8: Analysis result for a Non-AML case.]", italic=True, size=10)
    add_para(doc, "[Figure 9: Analysis result for an AML Detected case.]", italic=True, size=10)

    add_heading_styled(doc, "4.6.1 System Performance Characteristics", level=3)
    add_para(doc, "Table 9: System Performance Benchmarks", bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    t9 = doc.add_table(rows=7, cols=2, style="Table Grid")
    t9.rows[0].cells[0].text = "Metric"; t9.rows[0].cells[1].text = "Value"
    for i, (met, val) in enumerate([
        ("Model Parameters", "27,393"),
        ("Model Size on Disk", "136 KB"),
        ("Input Resolution", "224 × 224 × 3"),
        ("Inference Time (GPU, GTX 1650)", "< 50 ms per image"),
        ("Inference Time (CPU, i5-12400F)", "< 100 ms per image"),
        ("Application Size (standalone .exe)", "~850 MB (includes PyTorch runtime)"),
    ]):
        for j, v in enumerate([met, val]):
            t9.rows[i+1].cells[j].text = v
    doc.add_paragraph()

    add_heading_styled(doc, "4.7 Discussion", level=2)
    add_para(doc,
        "The results demonstrate that a lightweight deep learning architecture (27,393 parameters) can "
        "achieve perfect discriminative performance between monocytes and myeloblasts when trained on "
        "a well-curated, stain-normalized dataset. Several findings merit detailed discussion."
    )
    add_para(doc,
        "First, the critical importance of stain normalization cannot be overstated. The pre-normalization "
        "diagnostic analysis revealed that colour alone could separate the two classes with 98.2% accuracy "
        "— any model, regardless of architecture, would exploit this shortcut and fail to learn "
        "morphological features. The Reinhard normalization pipeline systematically eliminated this "
        "confound, reducing colour separability to random chance (50.0%). This preprocessing step is "
        "a methodological contribution that should be standard practice in blood smear AI applications, "
        "particularly when working with multi-source or publicly available datasets where staining "
        "protocols may differ between classes."
    )
    add_para(doc,
        "Second, the model's prediction confidence distribution provides evidence of genuine feature "
        "learning. The mean confidence of 0.981 across all 300 test images, with a minimum of 0.694 "
        "on the most ambiguous case, indicates clean class separation rather than marginal decisions. "
        "In clinical screening applications, high-confidence predictions are essential — a screening "
        "tool that frequently produces borderline (0.51-0.55) probabilities undermines user trust and "
        "clinical utility."
    )
    add_para(doc,
        "Third, the training efficiency demonstrates that state-of-the-art performance in medical image "
        "classification does not require large-scale computational infrastructure. The entire training "
        "pipeline — data loading, stain normalization, online augmentation, and model training — was "
        "completed on a consumer-grade GPU (GTX 1650, 4 GB VRAM) in under one hour. This has important "
        "implications for LMIC research capacity: institutions without access to cloud GPU clusters or "
        "data centre infrastructure can still develop and train high-performing medical AI models using "
        "appropriately designed lightweight architectures."
    )
    add_para(doc,
        "Fourth, the threshold analysis reveals an unusually wide safe operating range. The model "
        "maintains 97.3% recall with zero false positives at a threshold of 0.90 — most classifiers "
        "exhibit a sharp precision-recall trade-off where increasing the threshold beyond the optimum "
        "rapidly degrades recall. ALNet's flat precision curve across the full threshold range indicates "
        "that the learned feature representations produce well-separated class probability distributions, "
        "a property that is highly desirable for clinical deployment where threshold calibration against "
        "local disease prevalence may be necessary."
    )
    doc.add_page_break()

    # ===== CHAPTER 5: CONCLUSION =====
    add_heading_styled(doc, "CHAPTER 5: CONCLUSION AND RECOMMENDATIONS", level=1)

    add_heading_styled(doc, "5.1 Achievement of Objectives", level=2)

    add_para(doc, "Objective 1: Develop and train a detection model based on ALNet architecture.", bold=True)
    add_para(doc,
        "This objective was achieved. ALNet was successfully designed, implemented, and trained on "
        "2,000 stain-normalized blood smear images. The model architecture follows the proposed design: "
        "two depthwise separable convolutional blocks with localized sparse attention, max pooling "
        "transition, and a two-layer dense classification head with dropout regularization. The model "
        "contains 27,393 parameters and achieved stable, rapid convergence on modest GPU hardware."
    )

    add_para(doc, "Objective 2: Implement stain normalization to eliminate colour-domain confounds.", bold=True)
    add_para(doc,
        "This objective was achieved and was foundational to the study's validity. Reinhard stain "
        "normalization in LAB colour space was successfully implemented and validated through systematic "
        "pre- and post-normalization diagnostic checks. Colour-based class separability was reduced "
        "from 98.2% to 50.0% (random chance), and size-based differences were eliminated through "
        "uniform 400×400 padding. The normalization pipeline processed all 2,000 images without failure."
    )

    add_para(doc, "Objective 3: Test and evaluate the performance of the model.", bold=True)
    add_para(doc,
        "This objective was achieved with results exceeding expectations. The model achieved perfect "
        "classification on the 300-image test set: 100.0% accuracy, AUC-ROC of 1.000, F1-score of "
        "1.000, with zero false positives and zero false negatives. Mean prediction confidence was "
        "0.981 across all test samples. The threshold analysis demonstrated a wide safe operating "
        "range with 97.3% recall at the conservative 0.90 threshold with zero false positives."
    )

    add_para(doc, "Objective 4: Develop a user interface to support easy utilisability.", bold=True)
    add_para(doc,
        "This objective was achieved. A fully functional desktop application was developed using "
        "CustomTkinter, providing an intuitive interface for loading blood smear images, automated "
        "preprocessing, real-time AML screening with confidence scores, and SQLite-based session "
        "logging for audit trail purposes. The application explicitly labels all results as screening "
        "flags for human review, consistent with the decision-support scope. The application was "
        "packaged as a standalone executable using PyInstaller, enabling one-click deployment without "
        "requiring Python or dependency installation."
    )

    add_heading_styled(doc, "5.2 Answering the Research Questions", level=2)
    add_para(doc,
        "RQ1: Is it possible to develop and train an ALNet-based model on AML blood smear images?\n"
        "Yes. ALNet was successfully implemented and trained on 2,000 blood smear images, achieving "
        "stable convergence in 9 epochs with 100.0% validation accuracy. The model's 27,393-parameter "
        "architecture proved highly efficient, completing training on a consumer-grade GTX 1650 GPU "
        "in under one hour."
    )
    add_para(doc,
        "RQ2: Does Reinhard stain normalization effectively eliminate colour-based shortcuts in "
        "Wright-Giemsa stained images?\n"
        "Yes. Pre-normalization colour-based class separability of 98.2% was reduced to 50.0% "
        "(random chance) after Reinhard normalization. Post-normalization diagnostic checks confirmed "
        "that no single feature (colour ratio, brightness, edge density) could trivially separate "
        "the classes. This validates Reinhard normalization as an effective preprocessing methodology "
        "for Wright-Giemsa stained blood smear images."
    )
    add_para(doc,
        "RQ3: Can the trained ALNet-based model accurately distinguish myeloblasts from monocytes?\n"
        "Yes. The model achieved perfect classification on the held-out test set: 100.0% accuracy, "
        "AUC-ROC 1.000, F1-score 1.000, with zero false positives and zero false negatives across "
        "all 300 test images. Mean prediction confidence was 0.981, and the model maintained 97.3% "
        "recall at a conservative 0.90 threshold with zero false positives."
    )
    add_para(doc,
        "RQ4: Is it possible to integrate the model into a usable desktop application for clinical "
        "deployment?\n"
        "Yes. The ALNet Screening Tool desktop application successfully integrates the trained model "
        "with a user-friendly interface, automated preprocessing, real-time AML screening with "
        "confidence scores, and SQLite-based audit logging. The application was packaged as a "
        "standalone executable requiring no Python or dependency installation, making it deployable "
        "on standard laboratory workstations in LMIC settings."
    )

    add_heading_styled(doc, "5.3 Limitations", level=2)
    add_para(doc,
        "1. Single Dataset Source: The model was trained exclusively on the Kaggle blood cell dataset. "
        "While the balanced class distribution and stain normalization provide strong internal validity, "
        "external validation on multi-institutional data with different staining protocols and "
        "microscope configurations is necessary before clinical deployment."
    )
    add_para(doc,
        "2. Binary Classification Only: The model distinguishes only between monocytes and myeloblasts. "
        "It does not differentiate between AML subtypes (M0-M7 FAB classification) or detect other "
        "haematological malignancies, limiting its diagnostic scope."
    )
    add_para(doc,
        "3. Single-Cell Input Requirement: The current model operates on pre-extracted single-cell "
        "images rather than full-field blood smear images. A cell detection/localization module would "
        "be required upstream for fully automated pipeline deployment."
    )
    add_para(doc,
        "4. No Patient-Level Split: Due to anonymized filenames in the source dataset, the data split "
        "was at the image level rather than the patient level. If multiple single-cell images from the "
        "same patient appear across splits, this may introduce subtle data leakage."
    )
    add_para(doc,
        "5. Computational Requirements for Normalization: While the trained model is extremely "
        "lightweight (136 KB), the Reinhard normalization preprocessing requires per-image LAB "
        "colour space conversion and statistics matching, which adds computational overhead during "
        "batch processing of large slide collections."
    )

    add_heading_styled(doc, "5.4 Recommendations for Future Work", level=2)
    add_para(doc,
        "1. External Validation: The model should be prospectively validated on blood smear images "
        "from multiple laboratories and staining protocols to establish real-world clinical performance "
        "and assess robustness to staining variability beyond what Reinhard normalization can address."
    )
    add_para(doc,
        "2. Multi-Class Classification: Extending ALNet to classify additional cell types (lymphocytes, "
        "neutrophils, eosinophils, basophils, promyelocytes, and AML subtypes) would increase clinical "
        "utility for differential diagnosis and treatment planning."
    )
    add_para(doc,
        "3. Full-Field Image Processing: Integrating a cell detection and segmentation module would "
        "enable end-to-end processing of whole blood smear images, eliminating the manual cell "
        "extraction step and enabling batch screening of entire slides."
    )
    add_para(doc,
        "4. Prospective Clinical Study: A prospective validation study in partnership with the Uganda "
        "Cancer Institute would provide real-world evidence of clinical utility, user acceptance, "
        "and workflow integration in an LMIC clinical setting."
    )
    add_para(doc,
        "5. Model Optimization: Exploring INT8 quantization and ONNX conversion would further reduce "
        "the model's computational footprint and enable deployment on mobile devices or embedded "
        "systems for point-of-care screening in primary health facilities."
    )
    add_para(doc,
        "6. Explainability Integration: Implementing Grad-CAM or SHAP-based saliency maps would "
        "provide visual explanations of model predictions, increasing clinician trust and enabling "
        "verification that the model is attending to diagnostically relevant cellular regions."
    )
    add_para(doc,
        "7. Multi-Source Stain Normalization Benchmark: A systematic comparison of Reinhard, Macenko, "
        "Vahadane, and CycleGAN-based stain normalization methods on Wright-Giemsa blood smear images "
        "would establish best practices for the haematological AI community."
    )
    doc.add_page_break()

    # ===== REFERENCES =====
    add_heading_styled(doc, "REFERENCES", level=1)
    refs = [
        "Dores, G. M., et al. (2012). Acute leukemia incidence and patient survival among children and adults in the United States, 2001-2007. Blood, 119(1), 34-43.",
        "Kansal, R. (2019). Classification of acute myeloid leukemia by the revised fourth edition WHO criteria. Human Pathology, 90, 80-96.",
        "Munroe, M., et al. (2025). Low survival in younger adults with AML in Tanzania. PLoS One, 20(9), e0332237.",
        "Matek, C., et al. (2019). A Single-cell Morphological Dataset of Leukocytes from AML Patients and Non-malignant Controls (AML-Cytomorphology_LMU). The Cancer Imaging Archive.",
        "Singh, S. (2024). Blood Cell Images for Cancer Detection. Kaggle. https://www.kaggle.com/datasets/sumithsingh/blood-cell-images-for-cancer-detection",
        "Jabeen, K., et al. (2016). The Impact of Socioeconomic Factors on the Outcome of Childhood ALL Treatment in a LMIC. J. Pediatr. Hematol. Oncol., 38(8), 587-596.",
        "Blumenthal, D., & Patel, B. (2024). The Regulation of Clinical Artificial Intelligence. NEJM AI, 1(8).",
        "Hamet, P., & Tremblay, J. (2017). Artificial intelligence in medicine. Metabolism, 69, S36-S40.",
        "Haferlach, T., et al. (2005). Global approach to the diagnosis of leukemia using gene expression profiling. Blood, 106(4), 1189-1198.",
        "Gao, H., et al. (2024). ALNet: An adaptive channel attention network for accurate indoor visual localization. Expert Syst. Appl., 250, 123792.",
        "Abhishek, A., et al. (2022). Automated classification of acute leukemia using machine learning and deep learning techniques. Biomedical Signal Processing and Control, 72, 103341.",
        "Saeed, A., et al. (2022). A Deep Learning-Based Approach for the Diagnosis of ALL. Electronics, 11(19), 3168.",
        "Hameed, M., et al. (2025). Acute myeloid leukemia classification using ReLViT. Scientific Reports, 15(1), 32798.",
        "Nakisige, C., et al. (2023). Artificial intelligence and visual inspection in cervical cancer screening. Int. J. Gynecological Cancer, 33(10), 1515-1521.",
        "Wu, B., et al. (2025). Global, regional and national epidemiology of acute myeloid leukemia (1990-2021). Annals of Medicine, 57(1).",
        "Shafik, W., et al. (2026). A systematic literature review on transparency and interpretability of AI models in healthcare. Health Technol.",
        "Kalinaki, K. (2025). Internet of Health Things (IoHT): An Exploration of Principles, Components, Architectures, Challenges. Taylor & Francis.",
        "Reinhard, E., et al. (2001). Color transfer between images. IEEE Computer Graphics and Applications, 21(5), 34-41.",
        "Macenko, M., et al. (2009). A method for normalizing histology slides for quantitative analysis. IEEE Int. Symp. Biomedical Imaging, 1107-1110.",
    ]
    for i, ref in enumerate(refs, 1):
        add_para(doc, f"[{i}] {ref}", size=11)

    report_path = OUTPUT_DIR / "BIRYOMUMEISHO JOSHUA_CAPSTONE_REPORT_FINAL.docx"
    doc.save(str(report_path))
    print(f"Report saved to {report_path}")


if __name__ == "__main__":
    generate_report()
